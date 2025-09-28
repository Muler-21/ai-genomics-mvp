import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
from io import BytesIO
import openai
from docx import Document

# ---------------------------
# Sidebar: API key input
# ---------------------------
def set_openai_key_from_sidebar():
    st.sidebar.subheader("🔑 OpenAI API Key")
    api_key = st.sidebar.text_input(
        "Enter your OpenAI API key",
        type="password",
        help="Get a free key from https://platform.openai.com/"
    )
    if api_key:
        openai.api_key = api_key
        return api_key
    return None

api_key = set_openai_key_from_sidebar()
if not api_key:
    st.warning("⚠️ Please enter your OpenAI API key in the sidebar to continue.")
    st.info("👉 You can get a free key from OpenAI and paste it here. That way, **you pay nothing** as the app owner.")
    st.stop()

# ---------------------------
# Helpers
# ---------------------------
def extract_text_from_pdf(file):
    text = ""
    doc = fitz.open(stream=file.read(), filetype="pdf")
    for page in doc:
        text += page.get_text("text")
    return text

def ask_openai(prompt, model="gpt-4o-mini"):
    try:
        response = openai.ChatCompletion.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=3000,
        )
        return response["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return f"❌ Error: {e}"

def export_to_docx(text, title="AI Generated Report"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ---------------------------
# Streamlit App
# ---------------------------
st.set_page_config(page_title="AI Genomics Assistant", layout="wide")
st.title("🧬 AI Genomics & Literature Review Assistant")

# Welcome message
st.info("""
👋 **Welcome to the AI Genomics & Literature Review Assistant!**

Here’s how to use this app:

1. 👉 Open the **sidebar on the left** (click the `>` if hidden) and enter your **OpenAI API key**.  
   - Get one free at [platform.openai.com](https://platform.openai.com/).  
   - This ensures **you don’t pay anything as the app owner**.  

2. 📄 In the **Paper Summarizer** tab:  
   - Upload up to **10 PDFs or Word docs** (or paste text).  
   - Choose the detail level (**Concise**, **Expanded**, or **Very Detailed – up to 30–50 pages**).  
   - Click **Generate Coherent Report**.  
   - The app produces **100% humanized academic-style text** (summary, methodology, discussion, references).  
   - You can **copy-paste into Word** or **download as a .docx** file.  

3. 🧬 In the **Genomic Data Interpreter** tab:  
   - Upload a **VCF**, **CSV**, or **Excel** dataset.  
   - Preview your data instantly.  
   - Click **Interpret Genomic Data** for AI-generated insights.  
""")

tabs = st.tabs(["📄 Multi-Paper Summarizer", "🧬 Genomic Data Interpreter"])

# ---------------------------
# Tab 1: Multi-Paper Summarizer
# ---------------------------
with tabs[0]:
    st.header("Multi-Paper Summarizer & Coherent Report Builder")

    uploaded_files = st.file_uploader(
        "Upload up to 10 PDF or Word documents",
        type=["pdf", "txt", "docx"],
        accept_multiple_files=True
    )

    detail_level = st.radio(
        "Detail level",
        ["Concise", "Expanded", "Very Detailed (30–50 pages)"],
        horizontal=True
    )

    if st.button("Generate Coherent Report"):
        if not uploaded_files:
            st.error("❌ Please upload at least one file.")
        else:
            all_texts = []
            for file in uploaded_files[:10]:
                if file.type == "application/pdf":
                    all_texts.append(extract_text_from_pdf(file))
                elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    import docx
                    doc = docx.Document(file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    all_texts.append(text)
                else:
                    all_texts.append(file.read().decode("utf-8"))

            combined_text = "\n\n".join(all_texts)

            prompt = f"""
            You are an academic writer. Read the following collection of research papers and
            produce ONE SINGLE coherent long-form literature review.

            Requirements:
            - Write in natural, human-like academic style (never robotic).
            - Structure with: Abstract, Introduction, Methods (rewritten), Results, Discussion, Conclusion.
            - Ensure flow is coherent across papers, not just summaries.
            - Highlight similarities/differences between studies.
            - Add inline citations where appropriate.
            - End with a formatted References section.
            - Target length: {detail_level} (aim for 30–50 pages if 'Very Detailed').

            Texts to synthesize:
            {combined_text}
            """

            result = ask_openai(prompt)

            st.success("✅ Generated Coherent Report:")
            st.write(result)

            # Offer Word download
            docx_file = export_to_docx(result, "AI Generated Coherent Report")
            st.download_button(
                "⬇️ Download Full Report (Word)",
                data=docx_file,
                file_name="AI_Genomics_LitReview.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

# ---------------------------
# Tab 2: Genomic Data Interpreter
# ---------------------------
with tabs[1]:
    st.header("Upload Genomic Dataset (VCF / CSV / Excel)")
    gfile = st.file_uploader("Choose a file", type=["vcf","csv","xlsx"])

    if gfile is not None:
        df = None
        if gfile.name.endswith(".csv"):
            df = pd.read_csv(gfile)
        elif gfile.name.endswith(".xlsx"):
            df = pd.read_excel(gfile)
        elif gfile.name.endswith(".vcf"):
            content = gfile.read().decode("utf-8")
            rows = [line.split("\t") for line in content.splitlines() if not line.startswith("##")]
            df = pd.DataFrame(rows[1:], columns=rows[0])
        else:
            st.error("Unsupported file format")
            df = None

        if df is not None:
            st.write("📊 Preview of uploaded dataset:")
            st.dataframe(df.head())

            if st.button("Interpret Genomic Data"):
                prompt = f"""
                You are a genomics assistant. Interpret this dataset (first 20 rows shown):

                {df.head(20).to_csv(index=False)}

                Provide:
                - Summary of key findings
                - Possible biological/clinical interpretations
                - Limitations and cautions
                """
                result = ask_openai(prompt)
                st.success("✅ Genomic Interpretation:")
                st.write(result)

